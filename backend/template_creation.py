from docx import Document

def create_template():
    # Create a new Document
    doc = Document()

    # Add title and company information
    doc.add_heading('Consignment Collection Form', level=1)
    doc.add_paragraph('XYZ Pte Ltd')
    doc.add_paragraph('123 Random Street, Singapore, SG 123456')
    doc.add_paragraph('Phone: +65 1234 5678')
    doc.add_paragraph('Email: contact@xyzpteltd.com')
    doc.add_paragraph('\n')

    # Add placeholders for user to fill in their details
    doc.add_paragraph('Please fill out the following information to collect your consignment.')

    # Add fields with placeholders
    doc.add_paragraph('Name: {{ Name }}')
    doc.add_paragraph('Consignment ID: {{ Consignment_ID }}')
    doc.add_paragraph('Date of Collection: {{ Date }}')
    doc.add_paragraph('Contact Number: {{ Contact_Number }}')
    doc.add_paragraph('Signature: ____________________')
    
    # Save the template
    doc.save('Consignment_Collection_Template.docx')
    print("Template created: Consignment_Collection_Template.docx")

# Run the function to create the template
create_template()
